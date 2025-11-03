import os
from typing import List
import PyPDF2
from docx import Document
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDoc

class DocumentProcessor:
    def __init__(self, chunk_size=500, chunk_overlap=50):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
    
    def process_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                print(f"📄 Processing PDF: {total_pages} pages")
                
                for page_num in range(total_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    print(f"  ✓ Page {page_num + 1}/{total_pages} processed")
                
                if not text.strip():
                    print("⚠️  Warning: No text extracted from PDF")
                    print("⚠️  This might be a scanned/image-only PDF")
                    
        except Exception as e:
            print(f"❌ Error processing PDF: {e}")
            raise
        
        return text
    
    def process_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            print(f"📄 Processed DOCX: {len(doc.paragraphs)} paragraphs")
            return text
        except Exception as e:
            print(f"❌ Error processing DOCX: {e}")
            raise
    
    def process_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            print(f"📄 Processed TXT: {len(text)} characters")
            return text
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                print(f"📄 Processed TXT (latin-1): {len(text)} characters")
                return text
            except Exception as e:
                print(f"❌ Error processing TXT: {e}")
                raise
        except Exception as e:
            print(f"❌ Error processing TXT: {e}")
            raise
    
    def process_csv(self, file_path: str) -> str:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file_path)
            text = df.to_string()
            print(f"📊 Processed CSV: {len(df)} rows, {len(df.columns)} columns")
            return text
        except Exception as e:
            print(f"❌ Error processing CSV: {e}")
            raise
    
    def process_file(self, file_path: str, filename: str) -> List[LangchainDoc]:
        """Process file and return document chunks"""
        print(f"\n{'='*60}")
        print(f"📄 Processing: {filename}")
        print(f"{'='*60}")
        
        try:
            # Extract text based on file type
            if filename.lower().endswith('.pdf'):
                text = self.process_pdf(file_path)
            elif filename.lower().endswith('.docx'):
                text = self.process_docx(file_path)
            elif filename.lower().endswith('.txt'):
                text = self.process_txt(file_path)
            elif filename.lower().endswith('.csv'):
                text = self.process_csv(file_path)
            else:
                raise ValueError(f"Unsupported file type: {filename}")
            
            # Validate extracted text
            if not text or len(text.strip()) < 10:
                print(f"\n❌ ERROR: No meaningful text extracted")
                print(f"   File size: {os.path.getsize(file_path)} bytes")
                print(f"   Text length: {len(text)} characters")
                print(f"\n💡 Possible reasons:")
                print(f"   - Scanned PDF (image-only, no searchable text)")
                print(f"   - Corrupted file")
                print(f"   - Password-protected file")
                print(f"   - Empty document")
                raise ValueError(
                    "Could not extract text from document. "
                    "It might be a scanned PDF, image-only, or corrupted file. "
                    "Please try uploading a different version or a text-searchable PDF."
                )
            
            print(f"\n✅ Text extracted: {len(text)} characters")
            
            # Split into chunks
            print(f"🔄 Splitting text into chunks...")
            chunks = self.text_splitter.split_text(text)
            
            if not chunks or len(chunks) == 0:
                print(f"❌ ERROR: No chunks created from text")
                raise ValueError("Failed to split document into chunks")
            
            print(f"✅ Created {len(chunks)} chunks")
            
            # Create Document objects with metadata
            documents = [
                LangchainDoc(
                    page_content=chunk,
                    metadata={
                        "source": filename,
                        "chunk_index": i,
                        "total_chunks": len(chunks),
                        "file_type": filename.split('.')[-1].lower()
                    }
                )
                for i, chunk in enumerate(chunks)
            ]
            
            print(f"✅ Processing complete: {len(documents)} documents created")
            print(f"{'='*60}\n")
            
            return documents
            
        except ValueError as ve:
            # Re-raise validation errors
            print(f"{'='*60}\n")
            raise
        except Exception as e:
            print(f"\n❌ Unexpected error processing {filename}: {type(e).__name__}: {e}")
            print(f"{'='*60}\n")
            raise ValueError(f"Failed to process document: {str(e)}")